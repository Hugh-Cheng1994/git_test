import os
from engine.get_all_file import *
from aip import AipOcr
from docx import Document
from docx.shared import Inches


'''
脚本目的：实现图片转文字
实现方法：调用百度api
'''

# 得到文件夹下所有的图片文档
filePath = '..\\text_distinguish\\amage\\'
file_list = get_all_file_name(filePath)

# 调用百度api
""" 你的 APPID AK SK """
APP_ID = '24079464'
API_KEY = 'FjhGVRyMOuIjmv2m2pBo6Smc'
SECRET_KEY = 'cI6PTW9TgPuMVs4PjBG50gzT7heyutr1'

""" 读取图片 """
def get_file_content(filePath):
    with open(filePath, 'rb') as fp:
        return fp.read()

client = AipOcr(APP_ID, API_KEY, SECRET_KEY)

result_str = ''
for file_name in file_list:
    """ 调用通用文字识别, 图片参数为本地图片 """
    my_return = client.basicGeneral(get_file_content(filePath + file_name));
    word_list = my_return['words_result']
    for word in word_list:
        result_str += word['words']+'\n'

print(result_str)

document = Document()
document.add_heading('马俊英的文档', 0)
p = document.add_paragraph('以下是识别出来的文字'+'\n')
p.add_run(result_str)
document.save('mujunying.docx')



